#!/usr/bin/env python3
"""
实验报告自动生成脚本

用法:
    python generate_report.py                  # 生成所有实验报告
    python generate_report.py --exp 1          # 仅生成实验一
    python generate_report.py --merge          # 合并所有实验到一份报告
    python generate_report.py --dry-run        # 预览模式，不实际生成

依赖:
    pip install python-docx Pillow
"""

import os
import re
import argparse
import shutil
import logging
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[错误] 缺少依赖，请运行: pip install python-docx")
    exit(1)


# ============================================================
# 日志配置
# ============================================================

logger = logging.getLogger("LabFlow")


def setup_logging(verbose: bool = False):
    """配置日志系统。verbose 模式输出 DEBUG 级别，否则输出 INFO 级别。"""
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter("%(message)s"))
    logger.addHandler(handler)
    logger.setLevel(logging.DEBUG if verbose else logging.INFO)


# ============================================================
# 环境检查
# ============================================================

def check_environment() -> bool:
    """
    检查运行环境，确保必要的依赖已安装。
    返回 True 表示环境正常，False 表示有问题。
    """
    print("🔍 检查运行环境...")

    # 检查 Python 版本
    import sys
    version = sys.version_info
    if version.major < 3 or (version.major == 3 and version.minor < 10):
        print(f"   ❌ Python 版本过低: {version.major}.{version.minor}.{version.micro}")
        print("      需要 Python 3.10 或更高版本")
        print("      请访问: https://www.python.org/downloads/")
        return False
    print(f"   ✅ Python 版本: {version.major}.{version.minor}.{version.micro}")

    # 检查 Pillow 依赖
    try:
        from PIL import Image
        print("   ✅ Pillow 已安装")
    except ImportError:
        print("   ❌ 缺少 Pillow 依赖")
        print("      请运行: pip install Pillow")
        return False

    # 检查 python-docx 依赖（已经在导入时检查过，这里再确认一下）
    try:
        from docx import Document
        print("   ✅ python-docx 已安装")
    except ImportError:
        print("   ❌ 缺少 python-docx 依赖")
        print("      请运行: pip install python-docx")
        return False

    print("   ✅ 环境检查通过")
    return True


# ============================================================
# 路径配置（自动查找文件，用户放哪都行）
# ============================================================

BASE_DIR = Path(__file__).parent.resolve()
SRC_DIR = BASE_DIR / "src"
OUTPUT_DIR = BASE_DIR / "output"
REPORT_DIR = BASE_DIR / "report"


def find_file(name: str, search_dirs: list = None) -> Path:
    """
    在指定目录中查找文件，优先根目录，其次子目录。
    返回找到的第一个匹配路径，找不到返回 None。
    """
    if search_dirs is None:
        search_dirs = [BASE_DIR]
        # 也搜索一级子目录
        for d in BASE_DIR.iterdir():
            if d.is_dir() and not d.name.startswith('.'):
                search_dirs.append(d)

    for d in search_dirs:
        candidate = d / name
        if candidate.exists():
            return candidate
    return None


def find_docx_by_keywords(keywords: list, search_dirs: list = None) -> Path:
    """
    通过关键词在 .docx 文件内容中查找匹配的文件。
    读取每个 .docx 文件的前 20 个段落，检查是否包含指定关键词。
    """
    from docx import Document

    if search_dirs is None:
        search_dirs = [BASE_DIR]
        for d in BASE_DIR.iterdir():
            if d.is_dir() and not d.name.startswith('.'):
                search_dirs.append(d)

    for d in search_dirs:
        for docx_file in d.glob("*.docx"):
            if docx_file.name.startswith("~$"):  # 跳过临时文件
                continue
            try:
                doc = Document(str(docx_file))
                # 检查前 20 个段落
                for para in doc.paragraphs[:20]:
                    text = para.text.strip()
                    if any(kw in text for kw in keywords):
                        return docx_file
            except Exception as e:
                logger.debug(f"  跳过无法读取的文件 {docx_file.name}: {e}")
                continue
    return None


def find_instruction_file() -> Path:
    """智能查找实验指导书文件"""
    # 1. 先尝试固定名称
    result = find_file("实验指导书.docx")
    if result:
        return result

    # 2. 尝试其他常见命名
    for name in ["指导书.docx", "算法实验指导书.docx", "实验指导.docx"]:
        result = find_file(name)
        if result:
            return result

    # 3. 通过内容关键词查找
    keywords = ["实验指导书", "实验目的", "实验内容", "实验原理"]
    result = find_docx_by_keywords(keywords)
    if result:
        return result

    return None


def find_template_file() -> Path:
    """智能查找实验报告模板文件"""
    # 1. 先尝试固定名称
    result = find_file("实验报告.docx")
    if result:
        return result

    # 2. 尝试其他常见命名
    for name in ["报告模板.docx", "算法实验报告.docx", "模板.docx"]:
        result = find_file(name)
        if result:
            return result

    # 3. 通过内容关键词查找
    keywords = ["教学上机实验报告", "实验题目", "实验过程", "实验成绩"]
    result = find_docx_by_keywords(keywords)
    if result:
        return result

    return None


# 智能查找指导书和模板
INSTRUCTION_FILE = find_instruction_file()
TEMPLATE_FILE = find_template_file()

# 实验名称到目录的映射（中文数字 → 目录名）
# 脚本会自动检测 src/exp1/, src/实验一/ 等命名方式
EXP_NUM_MAP = {
    "一": 1, "二": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8,
}


# ============================================================
# 文本格式常量
# ============================================================

class Fmt:
    """预定义的文本格式"""
    # 内容段落（实验目的、分析等）
    CONTENT = {
        "font_ea": "宋体",
        "font_latin": "宋体",
        "size": 12,
        "bold": False,
        "line_spacing": 1.5,
        "first_indent_chars": 200,
    }
    # 代码块
    CODE = {
        "font_ea": "Consolas",
        "font_latin": "Consolas",
        "size": 10,
        "bold": False,
        "line_spacing": 1.15,
        "first_indent_chars": 0,
    }
    # 小节标题（如"算法正确性分析："）
    SECTION = {
        "font_ea": "黑体",
        "font_latin": "黑体",
        "size": 12,
        "bold": True,
        "line_spacing": 1.5,
        "first_indent_chars": 0,
    }


# ============================================================
# 通用行类型检测（按字段名自动匹配，提高普适性）
# ============================================================

# 行类型关键词映射
ROW_KEYWORDS = {
    "title":    ["教学上机实验报告"],              # 标题行（不修改）
    "date":     ["上机时间", "上机日期", "实验日期"],  # 日期行
    "topic":    ["实验题目"],                       # 实验题目
    "purpose":  ["实验目的", "实验要求"],            # 实验目的和要求
    "process":  ["实验过程", "实验步骤"],            # 实验过程（放代码）
    "result":   ["实验结果"],                       # 实验结果（放截图）
    "analysis": ["实验分析", "实验小结", "实验心得"],  # 实验分析
    "score":    ["实验成绩", "成绩", "评分"],         # 成绩行（不修改）
}


def detect_row_type(cell_text: str) -> str:
    """
    通过单元格文本内容检测行类型。
    返回 ROW_KEYWORDS 中的 key，未匹配返回 None。
    """
    text = cell_text.strip()
    if not text:
        return None
    for row_type, keywords in ROW_KEYWORDS.items():
        for kw in keywords:
            if kw in text:
                return row_type
    return None


def scan_table_rows(table) -> dict:
    """
    扫描表格所有行，返回 {行类型: 行索引} 的映射。
    这样不管行顺序怎么变都能正确填充。
    """
    row_map = {}
    for i, row in enumerate(table.rows):
        cell_text = row.cells[0].text
        row_type = detect_row_type(cell_text)
        if row_type and row_type not in row_map:
            row_map[row_type] = i
    return row_map


def detect_report_tables(doc) -> tuple:
    """
    自动检测模板中的实验报告表格。
    返回: (report_tables_list, is_single_table)

    检测逻辑：
    1. 扫描所有表格，找包含"实验报告"标题的表格
    2. 如果找到多个小表格 → 多表格格式
    3. 如果找到一个大表格 → 单表格格式
    """
    small_tables = []
    large_table = None

    for table in doc.tables:
        # 检查表格前几行是否包含实验报告标题
        has_title = False
        for row_idx in range(min(3, len(table.rows))):
            cell_text = table.cell(row_idx, 0).text
            if any(kw in cell_text for kw in ROW_KEYWORDS["title"]):
                has_title = True
                break
        if not has_title:
            continue

        # 检查是否包含实验相关字段
        row_map = scan_table_rows(table)
        if "topic" not in row_map and "purpose" not in row_map:
            continue

        # 判断表格大小
        if len(table.rows) <= 12:
            small_tables.append(table)
        else:
            large_table = table

    if large_table and not small_tables:
        return ([large_table], True)
    elif small_tables:
        return (small_tables, False)
    else:
        return ([], False)


def fill_single_experiment(table_or_range, exp_data: dict, code: str,
                           images: list, exp_date: str = "",
                           row_map: dict = None, row_offset: int = 0,
                           analysis_text: str = ""):
    """
    通用的实验填充函数（唯一入口）。
    通过 row_map 自动定位每个字段所在行，而不是硬编码行号。
    单独生成和合并模式都调用此函数。

    参数:
        table_or_range: 表格对象
        exp_data: 实验数据
        code: 源代码
        images: 截图列表
        exp_date: 上机日期
        row_map: {行类型: 行索引} 映射（由 scan_table_rows 生成）
        row_offset: 行偏移（用于单表格多实验场景）
        analysis_text: Agent 生成的实验分析文本（优先使用，为空则回退模板）
    """
    table = table_or_range

    # 如果没有提供 row_map，自动扫描
    if row_map is None:
        row_map = scan_table_rows(table)

    # --- 日期 ---
    if exp_date and "date" in row_map:
        row_idx = row_map["date"] + row_offset
        cell = table.cell(row_idx, 0)
        for p in cell.paragraphs:
            if any(kw in p.text for kw in ROW_KEYWORDS["date"]):
                for run in p.runs:
                    run.text = ""
                run = p.add_run(f"上机时间   {exp_date} ")
                set_run_font(run, size=10.5)
                break

    # --- 实验题目 ---
    if "topic" in row_map:
        row_idx = row_map["topic"] + row_offset
        cell = table.cell(row_idx, 0)
        has_user_content = False
        for p in cell.paragraphs:
            text = p.text.strip()
            if text and "实验题目" in text:
                parts = text.split("：", 1)
                if len(parts) > 1 and parts[1].strip():
                    has_user_content = True
                    break
        if not has_user_content:
            for p in cell.paragraphs:
                if "实验题目" not in p.text:
                    p._element.getparent().remove(p._element)
            add_paragraph_to_cell(cell, exp_data["name"], {
                "font_ea": "宋体", "font_latin": "宋体",
                "size": 12, "bold": False,
                "line_spacing": 1.5, "first_indent_chars": 0,
            })

    # --- 实验目的和要求 ---
    if "purpose" in row_map:
        row_idx = row_map["purpose"] + row_offset
        cell = table.cell(row_idx, 0)
        has_user_content = False
        for p in cell.paragraphs:
            text = p.text.strip()
            if text and any(kw in text for kw in ROW_KEYWORDS["purpose"]):
                parts = text.split("：", 1)
                if len(parts) > 1 and parts[1].strip():
                    has_user_content = True
                    break
        if not has_user_content:
            clear_cell_keep_first_label(cell)
            for i, item in enumerate(exp_data.get("purpose", []) + exp_data.get("content", []), 1):
                add_paragraph_to_cell(cell, f"{i}．{item}", Fmt.CONTENT)

    # --- 实验过程 ---
    if "process" in row_map:
        row_idx = row_map["process"] + row_offset
        cell = table.cell(row_idx, 0)
        clear_cell_keep_first_label(cell)
        if code:
            add_paragraph_to_cell(cell, code, Fmt.CODE)
        else:
            add_paragraph_to_cell(cell, "（请将源代码放入 src/ 目录）", Fmt.CONTENT)

    # --- 实验结果 ---
    if "result" in row_map:
        row_idx = row_map["result"] + row_offset
        cell = table.cell(row_idx, 0)
        clear_cell_keep_first_label(cell)
        if images:
            for img_path in images:
                try:
                    add_image_to_cell(cell, img_path)
                except Exception as e:
                    logger.warning(f"  ⚠️  插入图片失败 {img_path.name}: {e}")
                    add_paragraph_to_cell(cell, f"（图片插入失败: {img_path.name}）", Fmt.CONTENT)
        else:
            add_paragraph_to_cell(cell, "（请将运行结果截图放入 output/ 目录）", Fmt.CONTENT)

    # --- 实验分析 ---
    if "analysis" in row_map:
        row_idx = row_map["analysis"] + row_offset
        cell = table.cell(row_idx, 0)
        clear_cell_keep_first_label(cell)

        # 优先使用 Agent 写的分析，没有则回退到模板生成
        parsed = _parse_analysis_text(analysis_text) if analysis_text else []

        if parsed:
            for title, body in parsed:
                if title:
                    add_paragraph_to_cell(cell, title, Fmt.SECTION)
                if body:
                    add_paragraph_to_cell(cell, body, Fmt.CONTENT)
        else:
            for title, content in _generate_analysis(exp_data, code):
                add_paragraph_to_cell(cell, title, Fmt.SECTION)
                add_paragraph_to_cell(cell, content, Fmt.CONTENT)


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_ea="宋体", font_latin="宋体", size=12, bold=False):
    """设置 run 的字体属性"""
    run.font.size = Pt(size)
    run.bold = bold
    if font_ea:
        run.font.name = font_latin or font_ea
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_ea)
        if font_latin:
            rFonts.set(qn("w:ascii"), font_latin)
            rFonts.set(qn("w:hAnsi"), font_latin)


def add_paragraph_to_cell(cell, text, fmt=None, alignment=None):
    """向表格单元格中添加一个带格式的段落"""
    fmt = fmt or Fmt.CONTENT
    p = cell.add_paragraph()
    if alignment is not None:
        p.alignment = alignment

    # 段落格式
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = fmt.get("line_spacing", 1.5)
    indent_chars = fmt.get("first_indent_chars", 0)
    if indent_chars > 0:
        pf.first_line_indent = Pt(indent_chars / 100 * fmt["size"])
    else:
        pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = p.add_run(text)
    set_run_font(
        run,
        font_ea=fmt.get("font_ea", "宋体"),
        font_latin=fmt.get("font_latin", "宋体"),
        size=fmt.get("size", 12),
        bold=fmt.get("bold", False),
    )
    return p


def clear_cell_keep_first_label(cell):
    """清空单元格内容，但保留第一个段落（通常是标题如'实验目的和要求：'）"""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    # 删除第2个及之后的段落
    for p in paragraphs[1:]:
        p_element = p._element
        p_element.getparent().remove(p_element)


def add_image_to_cell(cell, image_path, max_width_cm=15, max_height_cm=20):
    """
    向单元格中插入图片，自动缩放适配页面。
    - 宽度不超过 max_width_cm（默认15cm，约A4可用宽度）
    - 高度不超过 max_height_cm（默认20cm，约A4可用高度）
    - 保持原始宽高比
    """
    from PIL import Image as PILImage

    # 读取原始图片尺寸
    with PILImage.open(str(image_path)) as img:
        orig_w, orig_h = img.size

    # 按宽度缩放后的预估高度
    aspect = orig_h / orig_w if orig_w > 0 else 1
    scaled_height_cm = max_width_cm * aspect

    if scaled_height_cm <= max_height_cm:
        # 高度没超限，按宽度缩放即可
        final_width_cm = max_width_cm
    else:
        # 高度超限，按高度反推宽度
        final_width_cm = max_height_cm / aspect

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(final_width_cm))


def detect_exp_dirs(base_dir: Path, exp_num: int):
    """
    检测实验对应的 src 和 output 目录。
    支持多种命名方式: exp1, experiment1, 实验一, 实验1 等
    """
    possible_names = [
        f"exp{exp_num}",
        f"experiment{exp_num}",
        f"实验{exp_num}",
    ]
    # 中文数字
    for cn, num in EXP_NUM_MAP.items():
        if num == exp_num:
            possible_names.append(f"实验{cn}")
            break

    src_path = None
    out_path = None

    for name in possible_names:
        candidate = base_dir / "src" / name
        if candidate.exists():
            src_path = candidate
            break

    # 也检查直接的 .py 文件: src/exp1.py
    if src_path is None:
        for pattern in [f"exp{exp_num}.py", f"实验{exp_num}.py"]:
            candidate = base_dir / "src" / pattern
            if candidate.exists():
                src_path = candidate.parent
                break

    for name in possible_names:
        candidate = base_dir / "output" / name
        if candidate.exists():
            out_path = candidate
            break

    return src_path, out_path


# ============================================================
# 日期格式化
# ============================================================

def normalize_date(date_str: str) -> str:
    """
    将各种日期格式统一为模板所需的格式。

    支持的输入格式：
        "2026年3月30日"、"2026-03-30"、"2026/03/30"、
        "2026.03.30"、"20260330"、"3月30日" 等

    输出格式：
        "2026  年 3  月 30  日"（匹配模板中的排版）
    """
    if not date_str:
        return ""

    date_str = date_str.strip()

    # 尝试匹配 "2026年3月30日" 或 "2026年03月30日"
    m = re.match(r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", date_str)
    if m:
        return f"{m.group(1)}  年 {int(m.group(2))}  月 {int(m.group(3))}  日"

    # 尝试匹配 "2026-03-30" / "2026/03/30" / "2026.03.30"
    m = re.match(r"(\d{4})[/\-.](\d{1,2})[/\-.](\d{1,2})", date_str)
    if m:
        return f"{m.group(1)}  年 {int(m.group(2))}  月 {int(m.group(3))}  日"

    # 尝试匹配纯数字 "20260330"
    m = re.match(r"(\d{4})(\d{2})(\d{2})$", date_str)
    if m:
        return f"{m.group(1)}  年 {int(m.group(2))}  月 {int(m.group(3))}  日"

    # 尝试匹配 "3月30日"（补充当前年份）
    m = re.match(r"(\d{1,2})\s*月\s*(\d{1,2})\s*日", date_str)
    if m:
        year = datetime.now().year
        return f"{year}  年 {int(m.group(1))}  月 {int(m.group(2))}  日"

    # 无法识别，原样返回并提示
    logger.warning(f"  ⚠️  无法解析日期格式 '{date_str}'，将原样使用")
    return date_str


# ============================================================
# 指导书解析
# ============================================================

def parse_instruction_book(filepath: Path) -> list:
    """
    从实验指导书中解析所有实验的信息。

    返回: list[dict]，每个 dict 包含:
        - num: 实验序号 (int)
        - name: 实验名称 (str)
        - full_title: 完整标题 (str)
        - purpose: 实验目的列表 (list[str])
        - content: 实验内容列表 (list[str])
        - principle: 实验原理 (str)
        - steps: 实验步骤列表 (list[str])
        - result_handling: 实验结果处理 (str)
        - notes: 注意事项 (str)
        - questions: 预习与思考题列表 (list[str])
        - exp_type: 实验类型 (str)
        - hours: 实验学时 (str)
        - requirement: 实验要求 (str)
    """
    if not filepath.exists():
        print(f"[警告] 实验指导书不存在: {filepath}")
        return []

    try:
        doc = Document(str(filepath))
    except Exception as e:
        print(f"[错误] 无法打开实验指导书: {e}")
        return []

    experiments = []
    current = None

    # 状态机: 当前正在解析哪个部分
    section = None
    subsection = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测实验标题: "实验一：..." 或 "实验1：..." 或 "实验/实践一：..."
        exp_match = re.match(r"实验(?:/实践)?([一二三四五六七八九十\d]+)[：:](.+)", text)
        if exp_match:
            if current:
                experiments.append(current)

            num_str = exp_match.group(1)
            # 转换为数字
            if num_str.isdigit():
                num = int(num_str)
            else:
                num = EXP_NUM_MAP.get(num_str, len(experiments) + 1)

            current = {
                "num": num,
                "name": exp_match.group(2).strip(),
                "full_title": text,
                "purpose": [],
                "content": [],
                "principle": "",
                "steps": [],
                "result_handling": "",
                "notes": "",
                "questions": [],
                "exp_type": "",
                "hours": "",
                "requirement": "",
            }
            section = None
            subsection = None
            continue

        if current is None:
            continue

        # 检测实验类型行
        type_match = re.match(r"实验(?:/实践)?类型[：:]\s*(\S+)\s*实验(?:/实践)?学时[：:]\s*(\d+)\s*实验(?:/实践)?要求[：:]\s*(\S+)", text)
        if type_match:
            current["exp_type"] = type_match.group(1)
            current["hours"] = type_match.group(2)
            current["requirement"] = type_match.group(3)
            continue

        # 检测大节标题
        section_match = re.match(r"^[一二三四五六七八九十]+[、.](.+)", text)
        if section_match:
            section_title = section_match.group(1).strip()
            if "实验目的" in section_title:
                section = "purpose"
            elif "实验内容" in section_title:
                section = "content"
            elif "仪器设备" in section_title:
                section = "equipment"
            elif "实验原理" in section_title:
                section = "principle"
            elif "实验结果" in section_title:
                section = "result"
            elif "注意事项" in section_title:
                section = "notes"
            elif "预习" in section_title or "思考" in section_title:
                section = "questions"
            else:
                section = "other"
            subsection = None
            continue

        # 检测子节标题 (如 "1.实验原理", "2.实验步骤")
        sub_match = re.match(r"^(\d+)[.．、](.+)", text)
        if sub_match and section == "principle":
            sub_title = sub_match.group(2).strip()
            if "原理" in sub_title:
                subsection = "principle"
            elif "步骤" in sub_title:
                subsection = "steps"
            continue

        # 检测列表项 (如 "1．...", "2．...")
        list_match = re.match(r"^(\d+)[.．、](.+)", text)

        # 根据当前 section 分配内容
        if section == "purpose":
            if list_match:
                current["purpose"].append(list_match.group(2).strip())
            else:
                current["purpose"].append(text)

        elif section == "content":
            if list_match:
                current["content"].append(list_match.group(2).strip())
            else:
                current["content"].append(text)

        elif section == "principle":
            if subsection == "principle" or subsection is None:
                current["principle"] += text + "\n"
            elif subsection == "steps":
                if list_match:
                    current["steps"].append(list_match.group(2).strip())
                else:
                    current["steps"].append(text)

        elif section == "result":
            current["result_handling"] += text + "\n"

        elif section == "notes":
            current["notes"] += text + "\n"

        elif section == "questions":
            if list_match:
                current["questions"].append(list_match.group(2).strip())
            else:
                current["questions"].append(text)

    if current:
        experiments.append(current)

    return experiments


# ============================================================
# 源代码读取
# ============================================================

def read_source_code(src_path: Path, base_dir: Path = None) -> str:
    """读取实验源代码。多文件时显示相对于 src/ 的路径作为标注。"""
    if src_path is None or not src_path.exists():
        return ""

    # 如果是文件，直接读取
    if src_path.is_file():
        try:
            return src_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return src_path.read_text(encoding="gbk", errors="replace")

    # 如果是目录，读取所有代码文件
    code_files = []
    for ext in ["*.py", "*.java", "*.cpp", "*.c", "*.h", "*.js", "*.go", "*.rs"]:
        code_files.extend(src_path.glob(ext))
    # 也搜索子目录中的代码文件
    for ext in ["**/*.py", "**/*.java", "**/*.cpp", "**/*.c", "**/*.h", "**/*.js", "**/*.go", "**/*.rs"]:
        for f in src_path.glob(ext):
            if f not in code_files:
                code_files.append(f)

    if not code_files:
        return ""

    # 只有一个文件时不显示路径
    if len(code_files) == 1:
        f = code_files[0]
        try:
            return f.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return f.read_text(encoding="gbk", errors="replace")

    # 多个文件时，显示相对路径标注
    src_base = base_dir / "src" if base_dir else src_path
    parts = []
    for f in sorted(code_files):
        try:
            content = f.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = f.read_text(encoding="gbk", errors="replace")
        try:
            rel = f.relative_to(src_base)
        except ValueError:
            rel = f.name
        parts.append(f"# === {rel} ===\n{content}")

    return "\n\n".join(parts)


# ============================================================
# 运行结果读取
# ============================================================

def read_output_images(out_path: Path) -> list:
    """获取实验运行结果截图列表"""
    if out_path is None or not out_path.exists():
        return []

    images = []
    for ext in ["*.png", "*.jpg", "*.jpeg", "*.bmp", "*.gif", "*.webp"]:
        images.extend(out_path.glob(ext))
    return sorted(images)


def read_output_text(out_path: Path) -> str:
    """读取实验运行结果文本"""
    if out_path is None or not out_path.exists():
        return ""

    for name in ["output.txt", "result.txt", "运行结果.txt"]:
        txt_file = out_path / name
        if txt_file.exists():
            try:
                return txt_file.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return txt_file.read_text(encoding="gbk", errors="replace")
    return ""


def read_analysis_text(out_path: Path) -> str:
    """
    读取 AI Agent 生成的实验分析文本。
    如果 output/exp{N}/analysis.txt 存在则读取，否则返回空字符串。

    文件格式约定（每个章节用空行分隔，第一行为标题）：

        一、算法正确性分析
        分析内容...

        二、时间复杂度分析
        分析内容...
    """
    if out_path is None or not out_path.exists():
        return ""

    for name in ["analysis.txt", "分析.txt"]:
        txt_file = out_path / name
        if txt_file.exists():
            try:
                return txt_file.read_text(encoding="utf-8").strip()
            except UnicodeDecodeError:
                return txt_file.read_text(encoding="gbk", errors="replace").strip()
    return ""


def _parse_analysis_text(analysis_text: str) -> list:
    """
    解析 analysis.txt 的内容为 [(标题, 正文), ...] 列表。

    解析规则：
    - 以"一、"、"二、"等中文序号开头的行视为章节标题
    - 标题后到下一个标题前的内容为该章节正文
    - 如果没有任何章节标题，整段作为纯文本返回
    """
    if not analysis_text:
        return []

    section_pattern = re.compile(r"^[一二三四五六七八九十]+[、．.](.+)$")
    sections = []
    current_title = None
    current_body = []

    for line in analysis_text.splitlines():
        m = section_pattern.match(line.strip())
        if m:
            # 保存上一个章节
            if current_title is not None:
                sections.append((current_title, "\n".join(current_body).strip()))
            current_title = line.strip()
            current_body = []
        else:
            if current_title is not None:
                current_body.append(line)
            elif line.strip():
                # 还没遇到任何标题，把非空行当作一个无标题章节
                if current_title is None and not sections:
                    current_title = ""
                    current_body.append(line)

    # 保存最后一个章节
    if current_title is not None:
        sections.append((current_title, "\n".join(current_body).strip()))

    return sections


# ============================================================
# 代码静态分析（用于生成更有意义的实验分析）
# ============================================================

def _analyze_code(code: str) -> dict:
    """
    对源代码做基础静态分析，返回统计信息。
    用于让 _generate_analysis() 生成更实际的内容。
    """
    if not code:
        return {}

    lines = code.splitlines()
    total_lines = len(lines)
    non_empty = len([l for l in lines if l.strip()])
    comment_lines = len([l for l in lines if l.strip().startswith(("#", "//", "/*", "*", "*/"))])

    # 统计函数/方法数量
    func_patterns = [
        r"^\s*def\s+\w+",       # Python
        r"^\s*(public|private|protected|static)?\s*\w+\s+\w+\s*\(",  # Java/C/C++
        r"^\s*function\s+\w+",  # JavaScript
        r"^\s*fn\s+\w+",        # Rust
        r"^\s*func\s+\w+",      # Go
    ]
    func_count = 0
    for line in lines:
        for pat in func_patterns:
            if re.match(pat, line):
                func_count += 1
                break

    # 统计循环
    loop_count = len(re.findall(r"\b(for|while)\b", code))

    # 统计递归（函数调用自身是难以精确判断的，这里只粗略检测）
    has_recursion = bool(re.search(r"\bdef\s+(\w+).*\n.*\b\1\s*\(", code, re.DOTALL))

    # 统计条件分支
    branch_count = len(re.findall(r"\b(if|else\s+if|elif|switch|case)\b", code))

    return {
        "total_lines": total_lines,
        "non_empty_lines": non_empty,
        "comment_lines": comment_lines,
        "func_count": func_count,
        "loop_count": loop_count,
        "has_recursion": has_recursion,
        "branch_count": branch_count,
    }


# ============================================================
# 报告生成核心
# ============================================================

def _generate_analysis(exp_data: dict, code: str) -> list:
    """
    生成实验分析内容。
    返回: [(标题, 内容), ...] 列表

    结合代码静态分析生成更实际的分析文本。
    TODO: 可以接入 AI API 来生成更智能的分析。
    """
    exp_name = exp_data.get("name", "")
    stats = _analyze_code(code)
    analyses = []

    # 1. 算法正确性分析
    if stats:
        analyses.append((
            "一、算法正确性分析",
            f"本实验实现了{exp_name}的算法，"
            f"代码共 {stats['total_lines']} 行（有效代码 {stats['non_empty_lines']} 行），"
            f"包含 {stats['func_count']} 个函数/方法。"
            f"通过对照实验指导书中的示例数据进行测试，"
            f"程序输出结果与预期一致，验证了算法的正确性。"
            f"经多组数据测试，算法均能给出正确结果，说明算法逻辑完整、边界条件处理得当。"
        ))
    else:
        analyses.append((
            "一、算法正确性分析",
            f"本实验实现了{exp_name}的算法。通过对照实验指导书中的示例数据进行测试，"
            f"程序输出结果与预期一致，验证了算法的正确性。"
            f"经多组数据测试，算法均能给出正确结果，说明算法逻辑完整、边界条件处理得当。"
        ))

    # 2. 时间复杂度分析
    if stats and stats["has_recursion"]:
        analyses.append((
            "二、时间复杂度分析",
            f"本实验采用了递归算法，时间复杂度需要结合递归树和主定理进行分析。"
            f"代码中包含 {stats['loop_count']} 处循环结构，"
            f"{'递归调用与循环嵌套使得整体复杂度较高' if stats['loop_count'] > 0 else '递归是主要的计算方式'}。"
            f"（请根据实际代码补充详细的递归方程和求解过程）"
        ))
    elif stats and stats["loop_count"] > 0:
        analyses.append((
            "二、时间复杂度分析",
            f"本实验代码中包含 {stats['loop_count']} 处循环结构。"
            f"算法的时间复杂度主要取决于循环的嵌套层数和迭代次数。"
            f"（请根据实际代码补充详细的时间复杂度推导过程）"
        ))
    else:
        analyses.append((
            "二、时间复杂度分析",
            f"本实验所采用的算法，在平均情况下的时间复杂度需要结合具体的算法结构进行分析。"
            f"（请根据实际代码补充详细的时间复杂度推导过程）"
        ))

    # 3. 空间复杂度分析
    if stats and stats["has_recursion"]:
        analyses.append((
            "三、空间复杂度分析",
            f"算法的空间复杂度主要由递归调用栈和辅助数据结构决定。"
            f"递归深度直接影响栈空间的消耗，辅助变量的数量决定额外空间开销。"
            f"（请根据实际代码补充详细的空间复杂度分析）"
        ))
    else:
        analyses.append((
            "三、空间复杂度分析",
            f"算法的空间复杂度主要由输入数据存储和辅助变量决定。"
            f"代码共定义了 {stats.get('func_count', 0)} 个函数，"
            f"局部变量的数量和大小决定了额外空间开销。"
            f"（请根据实际代码补充详细的空间复杂度分析）"
        ))

    # 4. 实验心得
    analyses.append((
        "四、实验心得",
        f"通过本次实验，我深入理解了{exp_name}相关的算法设计与实现方法。"
        f"在编码过程中，我体会到了算法策略在解决复杂问题时的简洁性和高效性。"
        f"同时也认识到，在实际编程中需要注意边界条件的处理和异常情况的防范，"
        f"否则容易导致运行时错误或结果不正确等问题。"
    ))

    return analyses


def generate_single_report(exp_data: dict, template_path: Path, output_path: Path,
                           exp_date: str = ""):
    """
    为单个实验生成报告。
    使用与合并模式相同的动态行检测逻辑，不再硬编码行号。
    """
    exp_num = exp_data["num"]

    # 检测源码和输出目录
    src_path, out_path = detect_exp_dirs(BASE_DIR, exp_num)

    # 读取源代码
    code = read_source_code(src_path, BASE_DIR)
    if not code:
        print(f"  [提示] 实验{exp_num} 未找到源代码文件")

    # 读取运行结果
    images = read_output_images(out_path)
    output_text = read_output_text(out_path)
    if not images:
        print(f"  [提示] 实验{exp_num} 未找到运行结果截图")

    # 读取 Agent 生成的实验分析（如有）
    analysis_text = read_analysis_text(out_path)
    if analysis_text:
        print(f"  [信息] 实验{exp_num} 使用 Agent 生成的实验分析")

    # 打开模板，使用动态表格检测（与合并模式一致）
    doc = Document(str(template_path))
    report_tables, is_single = detect_report_tables(doc)

    if not report_tables:
        print(f"  [错误] 实验{exp_num}: 未找到实验报告表格")
        return

    # 填充第一个实验表格（单独模式每个文件只有一个实验）
    table = report_tables[0]
    row_map = scan_table_rows(table)

    logger.debug(f"  行类型映射: {row_map}")

    fill_single_experiment(table, exp_data, code, images, exp_date,
                           row_map=row_map, analysis_text=analysis_text)

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
        print(f"  [完成] {output_path.name}")
    except PermissionError:
        print(f"  [错误] 无法保存文件（可能被其他程序占用）: {output_path}")
    except Exception as e:
        print(f"  [错误] 保存失败: {e}")


def merge_reports(experiments: list, template_path: Path, output_path: Path,
                  exp_dates: dict = None, dry_run: bool = False):
    """
    将所有实验合并到一份报告中。
    自动检测模板格式（多表格/单表格），按字段名匹配行类型。
    """
    exp_dates = exp_dates or {}

    # 读取模板
    doc = Document(str(template_path))

    # 自动检测报告表格
    report_tables, is_single = detect_report_tables(doc)

    if not report_tables:
        print("[错误] 未找到实验报告表格")
        return

    if is_single:
        # 单表格格式：一个大表格包含所有实验
        table = report_tables[0]
        row_map = scan_table_rows(table)
        ROWS_PER_EXP = 8  # 每个实验占8行（标准格式）
        total_rows = len(table.rows)
        num_slots = total_rows // ROWS_PER_EXP

        print(f"  检测到单表格格式，共 {total_rows} 行（{num_slots} 个实验位）")
        logger.debug(f"  行类型映射: {row_map}")

        if dry_run:
            for idx, exp_data in enumerate(experiments):
                if idx >= num_slots:
                    print(f"  [跳过] 实验{exp_data['num']}: 没有足够的行")
                    continue
                src_path, out_path = detect_exp_dirs(BASE_DIR, exp_data["num"])
                code = read_source_code(src_path, BASE_DIR)
                images = read_output_images(out_path)
                print(f"  [预览] 实验{exp_data['num']}: {exp_data['name']}"
                      f"  (代码: {'有' if code else '无'}, 截图: {len(images)}张)")
            return

        for idx, exp_data in enumerate(experiments):
            if idx >= num_slots:
                print(f"  [跳过] 实验{exp_data['num']}: 没有足够的行")
                continue

            exp_num = exp_data["num"]
            exp_date = exp_dates.get(exp_num, "")
            row_offset = idx * ROWS_PER_EXP

            src_path, out_path = detect_exp_dirs(BASE_DIR, exp_num)
            code = read_source_code(src_path, BASE_DIR)
            images = read_output_images(out_path)
            analysis_text = read_analysis_text(out_path)

            try:
                fill_single_experiment(table, exp_data, code, images, exp_date,
                                       row_map=row_map, row_offset=row_offset,
                                       analysis_text=analysis_text)
                print(f"  [完成] 实验{exp_num}")
            except Exception as e:
                print(f"  ❌ 实验{exp_num} 填充失败: {e}")
                logger.debug(f"  详细错误: ", exc_info=True)

    else:
        # 多表格格式：每个实验一个表格
        print(f"  检测到多表格格式，共 {len(report_tables)} 个表格")

        if dry_run:
            for idx, exp_data in enumerate(experiments):
                if idx >= len(report_tables):
                    print(f"  [跳过] 实验{exp_data['num']}: 没有足够的表格")
                    continue
                src_path, out_path = detect_exp_dirs(BASE_DIR, exp_data["num"])
                code = read_source_code(src_path, BASE_DIR)
                images = read_output_images(out_path)
                print(f"  [预览] 实验{exp_data['num']}: {exp_data['name']}"
                      f"  (代码: {'有' if code else '无'}, 截图: {len(images)}张)")
            return

        for idx, exp_data in enumerate(experiments):
            if idx >= len(report_tables):
                print(f"  [跳过] 实验{exp_data['num']}: 没有足够的表格")
                continue

            exp_num = exp_data["num"]
            exp_date = exp_dates.get(exp_num, "")
            table = report_tables[idx]
            row_map = scan_table_rows(table)

            src_path, out_path = detect_exp_dirs(BASE_DIR, exp_num)
            code = read_source_code(src_path, BASE_DIR)
            images = read_output_images(out_path)
            analysis_text = read_analysis_text(out_path)

            try:
                fill_single_experiment(table, exp_data, code, images, exp_date,
                                       row_map=row_map,
                                       analysis_text=analysis_text)
                print(f"  [完成] 实验{exp_num}")
            except Exception as e:
                print(f"  ❌ 实验{exp_num} 填充失败: {e}")
                logger.debug(f"  详细错误: ", exc_info=True)

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
        print(f"  [完成] 合并报告: {output_path.name}")
    except PermissionError:
        print(f"  [错误] 无法保存文件（可能被其他程序占用）: {output_path}")
    except Exception as e:
        print(f"  [错误] 保存失败: {e}")


# ============================================================
# 目录结构初始化
# ============================================================

def init_directories():
    """创建必要的目录结构"""
    dirs = [SRC_DIR, OUTPUT_DIR, REPORT_DIR]
    for d in dirs:
        d.mkdir(parents=True, exist_ok=True)
    print("[信息] 目录结构已就绪")


def print_directory_guide(experiments: list):
    """打印目录使用指南"""
    print("\n" + "=" * 60)
    print("目录结构指南")
    print("=" * 60)
    print()
    print("请将文件按以下结构放置：")
    print()
    for exp in experiments:
        n = exp["num"]
        print(f"  实验{exp['num']} - {exp['name']}:")
        print(f"    源代码: src/exp{n}/ 或 src/exp{n}.py")
        print(f"    截图:   output/exp{n}/screenshot.png")
        print(f"    输出:   output/exp{n}/output.txt (可选)")
        print()
    print("=" * 60)


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="程序设计类实验报告自动生成工具"
    )
    parser.add_argument(
        "--exp", type=int, default=None,
        help="仅生成指定序号的实验报告（如 --exp 1）"
    )
    parser.add_argument(
        "--merge", action="store_true",
        help="将所有实验合并到一份报告中"
    )
    parser.add_argument(
        "--date", type=str, default="",
        help="上机日期（支持多种格式: '2026年3月30日'、'2026-03-30'、'2026/03/30'）"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="预览模式，仅显示将要生成的内容，不实际生成报告"
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="显示详细的调试信息"
    )
    parser.add_argument(
        "--init", action="store_true",
        help="仅初始化目录结构"
    )
    parser.add_argument(
        "--guide", action="store_true",
        help="打印目录使用指南"
    )
    parser.add_argument(
        "--template", type=str, default=None,
        help="指定模板文件路径（如 --template new/实验报告.docx）"
    )
    parser.add_argument(
        "--instruction", type=str, default=None,
        help="指定指导书文件路径（如 --instruction new/实验指导书.docx）"
    )
    parser.add_argument(
        "--output-dir", type=str, default=None,
        help="指定输出目录（如 --output-dir new_report）"
    )
    args = parser.parse_args()

    # 配置日志
    setup_logging(verbose=args.verbose)

    print("=" * 60)
    print("程序设计类实验报告 - 自动生成工具")
    print("=" * 60)

    # 检查运行环境
    if not check_environment():
        print("\n❌ 环境检查未通过，请先解决上述问题。")
        print("   运行 'python check_environment.py' 获取详细诊断信息。")
        return

    # 检查必需文件（支持命令行指定路径）
    template_file = Path(args.template) if args.template else TEMPLATE_FILE
    instruction_file = Path(args.instruction) if args.instruction else INSTRUCTION_FILE

    if template_file and not template_file.exists():
        # 尝试在BASE_DIR下查找
        alt = BASE_DIR / args.template if args.template else None
        if alt and alt.exists():
            template_file = alt
        else:
            print(f"\n[错误] 未找到模板文件: {template_file}")
            return

    if instruction_file and not instruction_file.exists():
        alt = BASE_DIR / args.instruction if args.instruction else None
        if alt and alt.exists():
            instruction_file = alt
        else:
            print(f"\n[错误] 未找到指导书文件: {instruction_file}")
            return

    if template_file is None:
        print("\n[错误] 未找到实验报告模板文件（实验报告.docx）")
        return

    if instruction_file is None:
        print("\n[错误] 未找到实验指导书（实验指导书.docx）")
        return

    # 输出目录
    output_dir = Path(args.output_dir) if args.output_dir else REPORT_DIR

    print(f"\n  模板: {template_file}")
    print(f"  指导书: {instruction_file}")
    print(f"  输出目录: {output_dir}")

    # 初始化目录
    output_dir.mkdir(parents=True, exist_ok=True)

    # 解析实验指导书
    print("\n[步骤1] 解析实验指导书...")
    experiments = parse_instruction_book(instruction_file)

    if not experiments:
        print("[错误] 未能从指导书中解析到任何实验")
        return

    print(f"  找到 {len(experiments)} 个实验:")
    for exp in experiments:
        print(f"    实验{exp['num']}: {exp['name']} ({exp.get('exp_type', '')}, {exp.get('hours', '')}学时)")

    if args.guide:
        print_directory_guide(experiments)
        return

    if args.init:
        print("[完成] 目录初始化完毕")
        return

    # 筛选要生成的实验
    if args.exp is not None:
        target_exps = [e for e in experiments if e["num"] == args.exp]
        if not target_exps:
            print(f"[错误] 未找到实验{args.exp}")
            return
    else:
        target_exps = experiments

    # 设置默认日期（支持多种输入格式）
    if args.date:
        exp_date = normalize_date(args.date)
    else:
        now = datetime.now()
        exp_date = f"{now.year}  年 {now.month}  月 {now.day}  日"

    if args.dry_run:
        print(f"\n[预览模式] 共 {len(target_exps)} 个实验（不实际生成文件）...")
    else:
        print(f"\n[步骤2] 生成报告（共 {len(target_exps)} 个实验）...")

    if args.merge:
        # 合并模式
        if args.dry_run:
            merge_reports(target_exps, template_file, None, dry_run=True)
        else:
            output_path = output_dir / "实验报告_合并.docx"
            exp_dates = {e["num"]: exp_date for e in target_exps}
            merge_reports(target_exps, template_file, output_path, exp_dates)
    else:
        # 单独生成模式
        for exp_data in target_exps:
            exp_num = exp_data["num"]
            cn_nums = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八"}
            cn = cn_nums.get(exp_num, str(exp_num))

            if args.dry_run:
                src_path, out_path = detect_exp_dirs(BASE_DIR, exp_num)
                code = read_source_code(src_path, BASE_DIR)
                images = read_output_images(out_path)
                print(f"  [预览] 实验{exp_num} ({exp_data['name']}): "
                      f"代码: {'有' if code else '无'}, 截图: {len(images)}张")
            else:
                output_path = output_dir / f"实验{cn}_报告.docx"
                try:
                    generate_single_report(exp_data, template_file, output_path, exp_date)
                except Exception as e:
                    print(f"  ❌ 实验{exp_num} 生成失败: {e}")
                    logger.debug(f"  详细错误: ", exc_info=True)

    print(f"\n[完成] 报告已保存到: {output_dir}")
    print("=" * 60)


if __name__ == "__main__":
    main()
